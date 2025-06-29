import fitz
import json
import requests
import os
from io import BytesIO
from docx import Document
import streamlit as st

# === CONFIG ===
OPENROUTER_API_KEY = st.secrets.get("OPENROUTER_API_KEY") or os.getenv("OPENROUTER_API_KEY")
if not OPENROUTER_API_KEY:
    st.error("❌ OpenRouter API key not found. Add it in Streamlit > Settings > Secrets.")
    st.stop()

# === CLEAN UNICODE TEXT ===
def clean_text(text):
    return text.encode("utf-8", "ignore").decode("utf-8")

# === PDF TEXT EXTRACTION ===
def extract_pdf_text(uploaded_pdfs):
    combined_text = ""
    for file in uploaded_pdfs:
        with fitz.open(stream=file.read(), filetype="pdf") as doc:
            for page in doc:
                combined_text += page.get_text()
    return clean_text(combined_text)

# === PLACEHOLDER EXTRACTION FROM DOCX ===
def extract_placeholders(docx_file):
    doc = Document(docx_file)
    placeholders = set()
    for para in doc.paragraphs:
        for word in para.text.split():
            if word.startswith("[") and word.endswith("]"):
                placeholders.add(word.strip("[]"))
    return list(placeholders)

# === CALL LLM TO FILL PLACEHOLDERS ===
def call_llm(pdf_text, placeholders):
    prompt = f"""
You are an insurance report AI assistant. Extract values for the following placeholders from the report below:

{placeholders}

Text:
\"\"\"
{pdf_text[:6000]}
\"\"\"

Return ONLY valid JSON like:
{{"XM8_INSURED_NAME": "New Zion Hill Church", "XM8_DATE_INSPECTED": "2024-10-21", ...}}
"""
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://your-app-name.streamlit.app"  # optional but recommended
    }

    body = {
        "model": "mistralai/mixtral-8x7b",
        "messages": [{"role": "user", "content": prompt}]
    }

    try:
        encoded_body = json.dumps(body, ensure_ascii=False).encode("utf-8")
        res = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, data=encoded_body)
        res.raise_for_status()
        content = res.json()["choices"][0]["message"]["content"]
        return json.loads(content)
    except Exception as e:
        st.error(f"❌ LLM call failed: {e}")
        return {}

# === FALLBACK MOCK DATA ===
def mock_data():
    return {
        "XM8_INSURED_NAME": "NEW ZION HILL MISSIONARY BAPTIST CHURCH",
        "XM8_DATE_LOSS": "2024-10-21",
        "XM8_DATE_INSPECTED": "2024-10-21",
        "XM8_INSURED_P_STREET": "Unknown Street",
        "XM8_INSURED_P_CITY": "Unknown City",
        "XM8_INSURED_P_STATE": "TX",
        "XM8_INSURED_P_ZIP": "99999",
        "XM8_TOL_DESC": "Wind",
        "XM8_ESTIMATOR_NAME": "Steven Kujawski",
        "XM8_ESTIMATOR_E_MAIL": "commercialclaims@eberls.com",
        "XM8_ESTIMATOR_C_PHONE": "(800) 607-3604",
        "XM8_DATE_CURRENT": "2024-12-16"
    }

# === FILL TEMPLATE PLACEHOLDERS ===
def fill_template(docx_file, field_values):
    doc = Document(docx_file)
    for para in doc.paragraphs:
        for key, val in field_values.items():
            if f"[{key}]" in para.text:
                para.text = para.text.replace(f"[{key}]", val)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# === STREAMLIT APP ===
st.set_page_config("Eberl Report Auto-Filler", page_icon="📄")
st.title("📄 Eberl GuideOne Report Auto-Filler")

template_file = st.file_uploader("Upload GuideOne Template (.docx)", type=["docx"])
pdf_files = st.file_uploader("Upload Photo Reports (.pdf)", type=["pdf"], accept_multiple_files=True)

if st.button("Process Report"):
    if not template_file or not pdf_files:
        st.error("Please upload both the .docx template and at least one .pdf report.")
    else:
        with st.spinner("🔍 Extracting PDF text..."):
            text = extract_pdf_text(pdf_files)

        with st.spinner("🔎 Reading placeholders..."):
            placeholders = extract_placeholders(template_file)

        with st.spinner("🤖 Querying LLM..."):
            field_values = call_llm(text, placeholders)
            if not field_values:
                st.warning("⚠️ Using mock values due to LLM error.")
                field_values = mock_data()

        with st.spinner("📝 Filling template..."):
            filled_doc = fill_template(template_file, field_values)

        st.success("✅ Report ready!")
        st.download_button(
            label="📥 Download Filled Report",
            data=filled_doc,
            file_name="filled_eberl_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
